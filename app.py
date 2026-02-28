import requests
import uuid
import json
from io import BytesIO

from flask import Flask, render_template, request, jsonify, session, redirect, url_for, Response, send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth import auth_bp, login_required
from database import save_chat, get_chat_history

app = Flask(__name__)
app.secret_key = "scriptoria-secret-key"

@app.errorhandler(Exception)
def handle_exception(e):
    """Log any unhandled exception."""
    print(f"[SERVER ERROR] {str(e)}")
    import traceback
    traceback.print_exc()
    return jsonify({"error": str(e)}), 500

# Register Auth blueprint
app.register_blueprint(auth_bp)


# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def generate_title(prompt_text):
    """Use AI to generate a concise, thematic 3-5 word title for the story.

    Attempts to call the local Ollama model via `call_ollama`. Falls back to a
    deterministic heuristic if the call fails or returns nothing useful.
    """
    prompt = f"""You are a master screenwriter. 
Based ONLY on the following story idea, generate a short, punchy, cinematic title (3 to 5 words).
Do not include any quotes, preambles, or explanations. Just output the title.

Story Idea: {prompt_text}"""
    try:
        title = call_ollama(prompt).strip()
        title = title.replace('"', '').replace('`', '').replace('*', '')
        if not title:
            raise ValueError("Empty title")
        # Ensure it's not crazy long
        if len(title) > 60:
            title = title[:57] + "..."
        return title
    except Exception as e:
        print(f"[TITLE GEN ERROR] {e}. Falling back to deterministic.")
        # Fallback to simple heuristic
        words = prompt_text.split()
        meaningful = [w for w in words if len(w) > 2][:6]
        title = " ".join(meaningful)
        if len(title) > 40:
            title = title[:37] + "..."
        if not title:
            title = "Untitled Story"
        return title.capitalize()


# ─────────────────────────────────────────────────────────────
# Ollama helpers
# ─────────────────────────────────────────────────────────────

import time

def call_ollama_stream(prompt):
    """Call Ollama with streaming enabled – yields text chunks as they arrive."""
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "granite4:micro",
        "prompt": prompt,
        "stream": True
    }
    try:
        start_time = time.time()
        print(f"[OLLAMA] Request sent at {time.strftime('%H:%M:%S')}")
        response = requests.post(url, json=payload, timeout=180, stream=True)
        
        first_chunk = True
        for line in response.iter_lines():
            if line:
                if first_chunk:
                    print(f"[OLLAMA] First chunk received in {time.time() - start_time:.2f}s")
                    first_chunk = False
                try:
                    chunk = json.loads(line)
                    if "response" in chunk:
                        yield chunk["response"]
                except json.JSONDecodeError:
                    pass
        print(f"[OLLAMA] Total stream duration: {time.time() - start_time:.2f}s")
    except requests.exceptions.ConnectionError:
        yield "ERROR: Cannot connect to Ollama. Is it running at localhost:11434?"
    except Exception as e:
        yield f"ERROR: {str(e)}"

def call_ollama(prompt):
    """Call Ollama and return the full response string."""
    full_response = ""
    for chunk in call_ollama_stream(prompt):
        if chunk.startswith("ERROR:"):
            raise Exception(chunk)
        full_response += chunk
    return full_response


# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────


def translate_script_text(text, target_lang):
    """Translate a block of text using an external translation API.

    This helper currently uses the unofficial Google Translate JSON endpoint
    which does not require an API key. The format is simple and works for
    short/medium pieces of text; if you need an enterprise solution replace
    this logic with your preferred provider.

    The `target_lang` should be a two-letter ISO code (e.g. 'es', 'fr', 'hi').
    """
    if not text or not target_lang:
        return ""

    try:
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            "client": "gtx",
            "sl": "auto",
            "tl": target_lang,
            "dt": "t",
            "q": text
        }
        resp = requests.get(url, params=params, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        # result[0] is a list of translated segments
        translated = "".join([seg[0] for seg in result[0] if seg and seg[0]])
        return translated
    except Exception as err:
        raise Exception(f"Translation failed: {err}")


# ─────────────────────────────────────────────────────────────
# Core Routes
# ─────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/dashboard')
def dashboard():
    if "username" not in session:
        return redirect(url_for('index'))
    return render_template("dashboard.html", username=session["username"])


# ─────────────────────────────────────────────────────────────
# Story Generation  (Optimized Streaming)
# ─────────────────────────────────────────────────────────────

@app.route('/generate_story', methods=['POST'])
def generate_story():
    total_start = time.time()
    try:
        data = request.get_json(silent=True) or {}
        storyline = data.get("storyline", "").strip()
        provided_title = data.get("title", "").strip()
        char_ids  = data.get("character_ids", [])
        location  = data.get("location", "").strip()
        bgm_tone  = data.get("bgm", "").strip()

        if not storyline:
            return jsonify({"error": "Storyline missing"}), 400

        user_id_fixed = session.get("user_id")
        
        # --- CONTEXT INJECTION (Character Profiles) ---
        character_blobs = []
        if user_id_fixed and char_ids:
            from database import _get_client, DB_MODE
            if DB_MODE == "local":
                import local_db
                for cid in char_ids:
                    res = local_db._run_query("SELECT * FROM characters WHERE id = ?", (cid,), fetch_one=True)
                    if res: character_blobs.append(f"{res['name']}: {res['description']}. (Personality: {res.get('personality') or 'N/A'})")
            else:
                res = _get_client().table("characters").select("*").in_("id", char_ids).execute()
                for char in res.data:
                    character_blobs.append(f"{char['name']}: {char['description']}. (Personality: {char.get('personality') or 'N/A'})")
        
        char_context = "\n".join([f"- {blob}" for blob in character_blobs])
        
        # --- CINEMATIC PROMPT ---
        prompt = f"""You are a master cinematic screenwriter.
Write a short screenplay based on the story idea below.

SCENE SETTING:
{location if location else "Standard cinematic backdrop."}

CHARACTER PROFILES (Ensure these characters match these descriptions exactly):
{char_context if char_context else "No specific characters defined."}

STORY IDEA:
{storyline}

MUSICAL TONE:
{bgm_tone if bgm_tone else "Standard cinematic score."}

CRITICAL INSTRUCTIONS:
1. STRICT GENRE ADHERENCE: Write the screenplay strictly based on the provided STORY IDEA and SCENE SETTING. 
2. NO HALLUCINATION: DO NOT introduce unrelated historical events, characters, or locations (e.g., do not mention Indian independence, 1920s Delhi, or historical figures unless explicitly requested in the prompt).
3. Include [BGM: description] tags where the music should shift or emphasize the mood.
4. Output ONLY the screenplay text. No preambles.
"""

        print(f"\n[STORY START] user: {user_id_fixed} chars: {len(char_ids)} loc: {len(location)}")

        def generate():
            full_text = ""
            for chunk in call_ollama_stream(prompt):
                full_text += chunk
                yield f"data: {json.dumps({'text': chunk})}\n\n"

            # After generation, save to database and send metadata
            if user_id_fixed and full_text and not full_text.startswith("ERROR:"):
                try:
                    # Use provided title if available, otherwise ask AI
                    if provided_title:
                        title = provided_title
                    else:
                        title = generate_title(storyline)
                    # Convert char_ids to JSON string for storage
                    char_ids_json = json.dumps(char_ids)
                    item = save_chat(user_id_fixed, storyline, full_text, 
                                     title=title, 
                                     location=location, 
                                     bgm=bgm_tone, 
                                     char_ids=char_ids_json)
                    if item:
                        yield f"data: {json.dumps({'metadata': item})}\n\n"
                except Exception as db_err:
                    print(f"[DB ERROR] Save failed: {db_err}")
            
            print(f"[STORY COMPLETE] Total Time: {time.time() - total_start:.2f}s\n")

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        print(f"[ERROR] Generate failed: {str(e)}")
        return jsonify({"error": str(e)}), 500


# ─────────────────────────────────────────────────────────────
# History
# ─────────────────────────────────────────────────────────────

@app.route('/translate_script', methods=['POST'])
@login_required

def translate_script():
    """Accepts screenplay text and a target language, returns translated text.

    Request JSON should contain:
      - script: original screenplay string
      - target_language: two-letter ISO language code (e.g. "es", "fr")

    Response JSON:
      { "translated": "..." }
    """
    data = request.get_json(silent=True) or {}
    script_text = (data.get("script") or "").strip()
    target_lang = (data.get("target_language") or "").strip()

    if not script_text:
        return jsonify({"error": "Script text is required"}), 400
    if not target_lang:
        return jsonify({"error": "Target language is required"}), 400

    try:
        translated = translate_script_text(script_text, target_lang)
        return jsonify({"translated": translated}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/history', methods=['GET'])
@login_required
def history():
    """Return past generations (Optimized retrieval)."""
    user_id = request.current_user_id
    limit   = int(request.args.get("limit", 20))
    # Note: get_chat_history already uses limit(N) from previous fix
    entries = get_chat_history(user_id, limit=limit)
    return jsonify({"history": entries}), 200


@app.route('/rename_chat', methods=['POST'])
@login_required
def rename_chat():
    """Rename a specific item."""
    from database import update_chat_title
    data = request.json or {}
    chat_id = data.get("id")
    new_title = (data.get("title") or "").strip()

    if not chat_id or not new_title:
        return jsonify({"error": "Chat ID and title are required"}), 400

    result = update_chat_title(chat_id, request.current_user_id, new_title)
    if result:
        return jsonify({"success": True, "title": result["title"]}), 200
    return jsonify({"error": "Failed to rename"}), 500


# ─────────────────────────────────────────────────────────────
# Character Bible
# ─────────────────────────────────────────────────────────────

@app.route('/generate_cinematic_setting', methods=['POST'])
def generate_cinematic_setting():
    """Suggests an appropriate location or BGM tone based on the storyline."""
    try:
        data = request.get_json(silent=True) or {}
        storyline = data.get("storyline", "").strip()
        setting_type = data.get("type", "location")  # "location" or "bgm"

        if not storyline:
            return jsonify({"error": "Please provide a story idea first."}), 400

        type_label = "SCENE LOCATION / SCENERY" if setting_type == "location" else "BGM / MUSICAL TONE"

        prompt = f"""You are a creative screenplay consultant.
Based ONLY on the STORY IDEA below, suggest a concise, rich, and fitting {type_label}.

STORY IDEA:
{storyline}

REQUIREMENTS:
1. Return ONLY the description, nothing else. 
2. Keep it under 15 words.
3. Be highly specific to the genre, era, and mood implied by the STORY IDEA. (If it's fantasy, suggest fantasy elements. If historical, suggest historical).
4. DO NOT reference any events or locations not implied by the STORY IDEA.
"""

        suggestion = call_ollama(prompt).strip()
        # Clean up any potential markdown or quotes
        suggestion = suggestion.replace('"', '').replace('', '').strip()
        
        return jsonify({"suggestion": suggestion}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/suggest_title', methods=['POST'])
@login_required
def suggest_title():
    """Return an AI‑generated title for a provided story idea."""
    data = request.get_json(silent=True) or {}
    storyline = (data.get("storyline") or "").strip()
    if not storyline:
        return jsonify({"error": "Storyline is required"}), 400
    try:
        title = generate_title(storyline)
        return jsonify({"title": title}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/generate_character', methods=['POST'])
@login_required
def gen_character():
    """Use AI to generate a character profile based on a name or small prompt."""
    data = request.json or {}
    input_text = data.get("name", "").strip() or "a random unique character"
    storyline = data.get("storyline", "").strip()
    
    context_instruction = ""
    if storyline:
        context_instruction = f"""
STORY CONTEXT:
"{storyline}"

CRITICAL: The character MUST fit seamlessly into the above STORY CONTEXT. If the input name (like 'Kala Bhairava') belongs to a real-world or mythological figure, describe a character in THIS story who happens to have that name matching the genre of the STORY CONTEXT.
"""
    else:
        context_instruction = f"""
No specific story context provided. If the input name belongs to a historical or mythological figure (like 'Kala Bhairava'), describe them according to their actual legend or historical reality.
"""

    prompt = f"""You are a master cinematic storyteller and character writer.
Generate a detailed character profile for EXACTLY this name: "{input_text}"
{context_instruction}

CRITICAL INSTRUCTIONS:
1. YOU MUST USE THE EXACT NAME PROVIDED: "{input_text}". Do not swap it for a related movie character or reincarnation (e.g. do not change Kala Bhairava to Harsha).
2. Adapt to the genre implied by the input or STORY CONTEXT.
3. Keep the character cinematic, compelling, and grounded.

FORMAT YOUR RESPONSE EXACTLY LIKE THIS:
Name: [Character Name]
Description: [2-3 sentences on appearance, background, and role]
Personality: [Key traits, core values, and habits]
"""
    
    try:
        full_response = call_ollama(prompt)
        
        # Simple parser for the AI response
        lines = full_response.split("\n")
        char_data = {"name": "", "description": "", "personality": ""}
        for line in lines:
            if line.startswith("Name:"): char_data["name"] = line.replace("Name:", "").strip()
            if line.startswith("Description:"): char_data["description"] = line.replace("Description:", "").strip()
            if line.startswith("Personality:"): char_data["personality"] = line.replace("Personality:", "").strip()
        
        # Fallback if parsing fails
        if not char_data["name"] or not char_data["description"]:
            char_data["name"] = input_text if input_text != "a random unique character" else "Unknown Adventurer"
            char_data["description"] = full_response[:300]
            
        return jsonify(char_data), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/get_characters', methods=['GET'])
@login_required
def get_chars():
    """Retrieve all character profiles for the user."""
    from database import get_characters
    chars = get_characters(request.current_user_id)
    return jsonify({"characters": chars}), 200


@app.route('/save_character', methods=['POST'])
@login_required
def save_char():
    """Save a character profile."""
    from database import save_character
    data = request.json or {}
    name = data.get("name", "").strip()
    desc = data.get("description", "").strip()
    pers = data.get("personality", "").strip()

    if not name or not desc:
        return jsonify({"error": "Name and description are required"}), 400

    result = save_character(request.current_user_id, name, desc, pers)
    if isinstance(result, dict):
        return jsonify({"success": True, "character": result}), 200
    
    # Return specific DB error message
    return jsonify({"error": result or "Failed to save character"}), 500


@app.route('/delete_character', methods=['POST'])
@login_required
def delete_char():
    """Delete a character profile."""
    from database import delete_character
    data = request.json or {}
    char_id = data.get("id")

    if not char_id:
        return jsonify({"error": "Character ID required"}), 400

    delete_character(char_id, request.current_user_id)
    return jsonify({"success": True}), 200

@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    data = request.json or {}
    screenplay = data.get('screenplay', '')
    location = data.get('location', '')
    bgm = data.get('bgm', '')
    char_ids = data.get('character_ids', [])
    if not screenplay: return jsonify({"error": "No content"}), 400

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # --- ADD CINEMATIC SETTINGS TO PDF ---
    if location or bgm:
        elements.append(Paragraph("CINEMATIC SETTINGS", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        if location:
            elements.append(Paragraph(f"<b>Location:</b> {location}", styles["Normal"]))
        if bgm:
            elements.append(Paragraph(f"<b>BGM Tone:</b> {bgm}", styles["Normal"]))
        elements.append(Spacer(1, 24))

    # --- ADD CHARACTER PROFILES TO PDF ---
    if char_ids:
        from database import _get_client, DB_MODE
        elements.append(Paragraph("CHARACTER PROFILES", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        
        char_blobs = []
        if DB_MODE == "local":
            import local_db
            for cid in char_ids:
                char = local_db._run_query("SELECT * FROM characters WHERE id = ?", (cid,), fetch_one=True)
                if char: char_blobs.append(char)
        else:
            res = _get_client().table("characters").select("*").in_("id", char_ids).execute()
            char_blobs = res.data
            
        for char in char_blobs:
            elements.append(Paragraph(f"<b>{char['name']}</b>", styles["Heading3"]))
            elements.append(Paragraph(char['description'], styles["Normal"]))
            if char.get('personality'):
                elements.append(Paragraph(f"<i>Personality: {char['personality']}</i>", styles["Italic"]))
            elements.append(Spacer(1, 12))
        
        elements.append(Paragraph("<hr/>", styles["Normal"])) # Separator
        elements.append(Spacer(1, 24))

    # Screenplay Content
    elements.append(Paragraph("SCREENPLAY", styles["Heading1"]))
    elements.append(Spacer(1, 12))
    for line in screenplay.split("\n"):
        safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        elements.append(Paragraph(safe or "&nbsp;", styles["Normal"]))
        elements.append(Spacer(1, 4))
    
    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='screenplay.pdf', mimetype='application/pdf')


@app.route('/download_docx', methods=['POST'])
def download_docx():
    data = request.json or {}
    screenplay = data.get('screenplay', '')
    location = data.get('location', '')
    bgm = data.get('bgm', '')
    char_ids = data.get('character_ids', [])
    if not screenplay: return jsonify({"error": "No content"}), 400

    doc = Document()
    for section in doc.sections:
        section.left_margin, section.right_margin = Inches(1.5), Inches(1)
        section.top_margin, section.bottom_margin = Inches(1), Inches(1)

    # --- ADD CINEMATIC SETTINGS TO DOCX ---
    if location or bgm:
        doc.add_heading('CINEMATIC SETTINGS', level=1)
        if location:
            p = doc.add_paragraph()
            p.add_run('Location: ').bold = True
            p.add_run(location)
        if bgm:
            p = doc.add_paragraph()
            p.add_run('BGM Tone: ').bold = True
            p.add_run(bgm)
        doc.add_paragraph() # Spacer
        
    # --- ADD CHARACTER PROFILES TO DOCX ---
    if char_ids:
        from database import _get_client, DB_MODE
        doc.add_heading('CHARACTER PROFILES', level=1)
        
        char_blobs = []
        if DB_MODE == "local":
            import local_db
            for cid in char_ids:
                char = local_db._run_query("SELECT * FROM characters WHERE id = ?", (cid,), fetch_one=True)
                if char: char_blobs.append(char)
        else:
            res = _get_client().table("characters").select("*").in_("id", char_ids).execute()
            char_blobs = res.data
            
        for char in char_blobs:
            p = doc.add_paragraph()
            p.add_run(char['name']).bold = True
            doc.add_paragraph(char['description'])
            if char.get('personality'):
                p_p = doc.add_paragraph()
                p_p.add_run(f"Personality: {char['personality']}").italic = True
            doc.add_paragraph() # Spacer
            
        doc.add_page_break()

    doc.add_heading('SCREENPLAY', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in screenplay.split("\n"):
        stripped = line.strip()
        para = doc.add_paragraph()
        if stripped.startswith(('INT.', 'EXT.', 'INT ', 'EXT ')) or (stripped.isupper() and 0 < len(stripped) < 60):
            para.add_run(stripped).bold = True
        else:
            para.add_run(stripped)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='screenplay.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


if __name__ == "__main__":
    app.run(debug=True)